from __future__ import annotations

import io
from datetime import datetime
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from api.deps.current_user import get_current_company
from database.database import get_db
from models.company import Company
from services.report_service import ReportService


router = APIRouter(
    prefix="/reports",
    tags=["Reports"],
)


def _report_payload(report):
    return {
        "id": str(report.id),
        "title": report.title,
        "client_name": (
            report.client.business_name
            if report.client
            else "N/A"
        ),
        "date": (
            report.generated_at.isoformat()
            if report.generated_at
            else None
        ),
        "score": float(report.score or 0),
        "format": report.format,
        "content": report.content or "",
        "summary": report.summary or "",
        "audit_id": (
            str(report.audit_id)
            if report.audit_id
            else None
        ),
    }


@router.get("/")
def get_reports(
    limit: int = Query(
        default=50,
        ge=1,
        le=200,
    ),
    db: Session = Depends(get_db),
    company: Company = Depends(get_current_company),
):
    service = ReportService(db)

    reports = service.get_reports_for_company(
        str(company.id),
        limit,
    )

    return [
        _report_payload(report)
        for report in reports
    ]


@router.get("/{report_id}")
def get_report(
    report_id: str,
    db: Session = Depends(get_db),
    company: Company = Depends(get_current_company),
):
    service = ReportService(db)

    report = service.get_report(
        report_id,
        str(company.id),
    )

    if not report:
        raise HTTPException(
            status_code=404,
            detail="Report not found",
        )

    return _report_payload(report)


@router.delete("/{report_id}")
def delete_report(
    report_id: str,
    db: Session = Depends(get_db),
    company: Company = Depends(get_current_company),
):
    service = ReportService(db)

    if not service.delete_report(
        report_id,
        str(company.id),
    ):
        raise HTTPException(
            status_code=404,
            detail="Report not found",
        )

    return {
        "success": True,
        "message": "Report deleted successfully.",
    }


def _build_docx(report) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX generation."
        ) from exc

    document = Document()

    title = document.add_heading(
        report.title,
        level=0,
    )

    title.runs[0].font.size = Pt(22)

    document.add_paragraph(
        f"Score: {float(report.score or 0):.1f}/100"
    )

    if report.generated_at:
        document.add_paragraph(
            "Generated: "
            + report.generated_at.strftime(
                "%d %B %Y %H:%M UTC"
            )
        )

    document.add_paragraph(
        report.summary or ""
    )

    content = report.content or ""

    for raw_line in content.splitlines():
        line = raw_line.rstrip()

        if not line:
            document.add_paragraph("")
            continue

        if line.startswith("# "):
            document.add_heading(
                line[2:].strip(),
                level=1,
            )
            continue

        if line.startswith("## "):
            document.add_heading(
                line[3:].strip(),
                level=2,
            )
            continue

        if line.startswith("### "):
            document.add_heading(
                line[4:].strip(),
                level=3,
            )
            continue

        if line.startswith("- "):
            document.add_paragraph(
                line[2:].strip(),
                style="List Bullet",
            )
            continue

        if line.startswith("> "):
            document.add_paragraph(
                line[2:].strip()
            )
            continue

        document.add_paragraph(
            line
        )

    output = io.BytesIO()

    document.save(output)

    return output.getvalue()


def _build_pdf(report) -> bytes:
    try:
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import (
            ParagraphStyle,
            getSampleStyleSheet,
        )
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except ImportError as exc:
        raise RuntimeError(
            "reportlab is required for PDF generation."
        ) from exc

    output = io.BytesIO()

    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontSize=20,
        leading=24,
        spaceAfter=10,
    )

    h1_style = ParagraphStyle(
        "ReportH1",
        parent=styles["Heading1"],
        fontSize=15,
        leading=19,
        spaceBefore=10,
        spaceAfter=6,
    )

    h2_style = ParagraphStyle(
        "ReportH2",
        parent=styles["Heading2"],
        fontSize=12,
        leading=15,
        spaceBefore=8,
        spaceAfter=5,
    )

    body_style = ParagraphStyle(
        "ReportBody",
        parent=styles["BodyText"],
        fontSize=9.5,
        leading=14,
        spaceAfter=5,
    )

    story = []

    story.append(
        Paragraph(
            report.title,
            title_style,
        )
    )

    story.append(
        Paragraph(
            f"Score: {float(report.score or 0):.1f}/100",
            body_style,
        )
    )

    if report.generated_at:
        story.append(
            Paragraph(
                "Generated: "
                + report.generated_at.strftime(
                    "%d %B %Y %H:%M UTC"
                ),
                body_style,
            )
        )

    story.append(
        Spacer(1, 8)
    )

    # Simple Markdown-to-PDF conversion.
    for raw_line in (report.content or "").splitlines():
        line = raw_line.strip()

        if not line:
            story.append(
                Spacer(1, 4)
            )
            continue

        escaped = (
            line.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

        if escaped.startswith("# "):
            story.append(
                Paragraph(
                    escaped[2:],
                    h1_style,
                )
            )

        elif escaped.startswith("## "):
            story.append(
                Paragraph(
                    escaped[3:],
                    h1_style,
                )
            )

        elif escaped.startswith("### "):
            story.append(
                Paragraph(
                    escaped[4:],
                    h2_style,
                )
            )

        elif escaped.startswith("- "):
            story.append(
                Paragraph(
                    "• " + escaped[2:],
                    body_style,
                )
            )

        elif escaped.startswith("> "):
            story.append(
                Paragraph(
                    escaped[2:],
                    body_style,
                )
            )

        else:
            story.append(
                Paragraph(
                    escaped,
                    body_style,
                )
            )

    document.build(story)

    return output.getvalue()


@router.get("/{report_id}/download")
def download_report(
    report_id: str,
    format: str = Query(
        default="pdf",
        pattern="^(pdf|docx)$",
    ),
    db: Session = Depends(get_db),
    company: Company = Depends(get_current_company),
):
    service = ReportService(db)

    report = service.get_report(
        report_id,
        str(company.id),
    )

    if not report:
        raise HTTPException(
            status_code=404,
            detail="Report not found",
        )

    try:
        if format == "docx":
            content = _build_docx(report)

            filename = (
                f"{report.title}.docx"
                .replace("/", "-")
                .replace("\\", "-")
            )

            media_type = (
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            )

        else:
            content = _build_pdf(report)

            filename = (
                f"{report.title}.pdf"
                .replace("/", "-")
                .replace("\\", "-")
            )

            media_type = "application/pdf"

    except RuntimeError as exc:
        raise HTTPException(
            status_code=500,
            detail=str(exc),
        ) from exc

    safe_filename = quote(
        filename
    )

    return StreamingResponse(
        io.BytesIO(content),
        media_type=media_type,
        headers={
            "Content-Disposition": (
                f"attachment; filename*=UTF-8''{safe_filename}"
            )
        },
    )